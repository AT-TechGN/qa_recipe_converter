import logging
import unicodedata
from docx import Document

logger = logging.getLogger(__name__)

# Mapping des colonnes reconnues (insensible à la casse, variantes incluses)
COLUMN_ALIASES = {
    'use_case_id': [
        'use case', 'usecase', 'uc', 'id', 'numero', 'n°', '#',
        'cas de test', 'cas', 'identifiant'
    ],
    'description': [
        'description', 'titre', 'intitule', 'libelle',
        'objectif', 'nom du test', 'nom'
    ],
    'preconditions': [
        'preconditions', 'pre-conditions',
        'prerequis', 'contexte', 'donnees initiales', 'donnees'
    ],
    'steps': [
        'etapes', 'actions', 'scenario', 'deroulement',
        'procedure', 'etapes de test', 'actions a realiser'
    ],
    'expected_results': [
        'resultats attendus', 'resultat attendu',
        'attendu', 'expected', 'comportement attendu'
    ],
    'observed_results': [
        'resultats observes', 'resultat observe',
        'observe', 'observed', 'obtenu', 'resultat obtenu'
    ],
}


def _strip_accents(text: str) -> str:
    """Remove accents for fuzzy matching."""
    return ''.join(
        c for c in unicodedata.normalize('NFD', text)
        if unicodedata.category(c) != 'Mn'
    )


def normalize(text: str) -> str:
    """Normalize text: lowercase, strip, remove accents, collapse spaces."""
    text = text.strip().lower().replace('\n', ' ').replace('\r', '')
    return _strip_accents(text)


def detect_column(header: str) -> str | None:
    """Return the field name for a given header string, or None."""
    h = normalize(header)
    for field, aliases in COLUMN_ALIASES.items():
        for alias in aliases:
            if alias in h or h in alias:
                return field
    return None


class DocxParser:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = None

    def _load_document(self):
        try:
            self.document = Document(self.file_path)
        except Exception as e:
            raise ValueError(f"Impossible d'ouvrir le fichier Word : {e}")

    def extract_use_cases(self) -> list[dict]:
        """Extract all use cases from all tables in the document."""
        self._load_document()
        all_use_cases = []

        if not self.document.tables:
            logger.warning("Aucun tableau trouvé dans le document.")
            return []

        for table_idx, table in enumerate(self.document.tables):
            logger.info(f"Analyse du tableau {table_idx + 1}...")
            use_cases = self._parse_table(table, table_idx)
            all_use_cases.extend(use_cases)

        logger.info(f"Total : {len(all_use_cases)} use cases extraits.")
        return all_use_cases

    def _get_cell_text(self, cell) -> str:
        """Extract clean text from a cell, handling merged cells."""
        text = cell.text.strip()
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return '\n'.join(lines)

    def _parse_table(self, table, table_idx: int) -> list[dict]:
        """Parse a single table and return list of use case dicts."""
        rows = table.rows
        if len(rows) < 2:
            logger.debug(f"Tableau {table_idx + 1} ignoré (moins de 2 lignes).")
            return []

        # Detect header row (first row)
        header_row = rows[0]
        column_map = {}  # col_index -> field_name

        for col_idx, cell in enumerate(header_row.cells):
            cell_text = self._get_cell_text(cell)
            field = detect_column(cell_text)
            if field and field not in column_map.values():
                column_map[col_idx] = field
                logger.debug(f"Colonne détectée : '{cell_text}' -> '{field}'")

        if not column_map:
            logger.warning(f"Tableau {table_idx + 1} : aucune colonne reconnue, ignoré.")
            return []

        logger.info(f"Tableau {table_idx + 1} : colonnes = {list(column_map.values())}")

        use_cases = []
        last_values = {}

        for row_idx, row in enumerate(rows[1:], start=1):
            uc_data = {
                'use_case_id': '',
                'description': '',
                'preconditions': '',
                'steps': '',
                'expected_results': '',
                'observed_results': '',
            }

            row_is_empty = True

            for col_idx, field_name in column_map.items():
                if col_idx >= len(row.cells):
                    continue

                cell = row.cells[col_idx]
                cell_text = self._get_cell_text(cell)

                if not cell_text and col_idx in last_values:
                    cell_text = last_values[col_idx]
                elif cell_text:
                    last_values[col_idx] = cell_text
                    row_is_empty = False

                uc_data[field_name] = cell_text

            if row_is_empty:
                continue

            # Skip rows that look like repeated header rows (all cells exactly match known headers)
            filled_vals = [v for v in uc_data.values() if v]
            if filled_vals and len(filled_vals) >= 3:
                # Only skip if ALL values are exact/very close header aliases (max 25 chars)
                def _is_header_word(v):
                    n = normalize(v)
                    return len(n) <= 25 and any(
                        n == alias for aliases in COLUMN_ALIASES.values() for alias in aliases
                    )
                if all(_is_header_word(v) for v in filled_vals):
                    continue

            use_cases.append(uc_data)

        return use_cases
