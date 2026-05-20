import os
import pytest
from docx import Document
from docx.shared import Pt
from apps.parser.docx_parser import DocxParser, detect_column, normalize, COLUMN_ALIASES


# --- Unit tests for helper functions ---

class TestNormalize:
    def test_strips_whitespace(self):
        assert normalize("  hello  ") == "hello"

    def test_lowercases(self):
        assert normalize("Description") == "description"

    def test_removes_newlines(self):
        assert normalize("étapes\nde test") == "etapes de test"


class TestDetectColumn:
    def test_detect_use_case(self):
        assert detect_column("Use Case") == "use_case_id"
        assert detect_column("UC") == "use_case_id"
        assert detect_column("use case") == "use_case_id"

    def test_detect_description(self):
        assert detect_column("Description") == "description"
        assert detect_column("DESCRIPTION") == "description"

    def test_detect_preconditions(self):
        assert detect_column("Préconditions") == "preconditions"
        assert detect_column("Preconditions") == "preconditions"
        assert detect_column("Pré-conditions") == "preconditions"

    def test_detect_steps(self):
        assert detect_column("Étapes") == "steps"
        assert detect_column("Actions") == "steps"
        assert detect_column("Scénario") == "steps"

    def test_detect_expected_results(self):
        assert detect_column("Résultats Attendus") == "expected_results"
        assert detect_column("Résultat Attendu") == "expected_results"
        assert detect_column("attendu") == "expected_results"

    def test_detect_observed_results(self):
        assert detect_column("Résultats Observés") == "observed_results"
        assert detect_column("Observé") == "observed_results"

    def test_unknown_column(self):
        assert detect_column("XYZ Unknown Column 123") is None


# --- Integration tests with real .docx files ---

def _create_test_docx(tmp_path, headers, rows):
    """Helper: create a .docx with a table."""
    doc = Document()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))

    # Header row
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    path = str(tmp_path / "test.docx")
    doc.save(path)
    return path


class TestDocxParser:
    def test_extracts_basic_use_cases(self, tmp_path):
        headers = ["Use Case", "Description", "Préconditions", "Étapes", "Résultats Attendus", "Résultats Observés"]
        rows = [
            ["UC001", "Login avec credentials valides", "Utilisateur existant", "1. Aller sur /login\n2. Saisir credentials", "Redirection dashboard", ""],
            ["UC002", "Login avec mauvais mot de passe", "Compte existant", "1. Saisir mauvais MDP", "Message d'erreur affiché", ""],
        ]
        path = _create_test_docx(tmp_path, headers, rows)
        parser = DocxParser(path)
        result = parser.extract_use_cases()

        assert len(result) == 2
        assert result[0]['use_case_id'] == 'UC001'
        assert result[0]['description'] == 'Login avec credentials valides'
        assert result[1]['use_case_id'] == 'UC002'

    def test_case_insensitive_headers(self, tmp_path):
        headers = ["USE CASE", "DESCRIPTION", "PRECONDITIONS", "ETAPES", "RESULTATS ATTENDUS", "RESULTATS OBSERVES"]
        rows = [["UC001", "Test", "Prérequis", "Action", "Résultat", ""]]
        path = _create_test_docx(tmp_path, headers, rows)
        parser = DocxParser(path)
        result = parser.extract_use_cases()

        assert len(result) == 1
        assert result[0]['use_case_id'] == 'UC001'

    def test_empty_document_returns_empty(self, tmp_path):
        doc = Document()
        path = str(tmp_path / "empty.docx")
        doc.save(path)
        parser = DocxParser(path)
        result = parser.extract_use_cases()
        assert result == []

    def test_document_without_matching_headers(self, tmp_path):
        headers = ["Col A", "Col B", "Col C"]
        rows = [["val1", "val2", "val3"]]
        path = _create_test_docx(tmp_path, headers, rows)
        parser = DocxParser(path)
        result = parser.extract_use_cases()
        # No recognized columns → empty result
        assert result == []

    def test_multiple_tables(self, tmp_path):
        doc = Document()
        headers = ["Use Case", "Description", "Préconditions", "Étapes", "Résultats Attendus", "Résultats Observés"]

        for i in range(2):
            table = doc.add_table(rows=2, cols=len(headers))
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h
            row_data = [f"UC{i+1}0{j+1}", f"Desc {i}", "Prérequis", "Actions", "Attendu", ""]
            for j, val in enumerate(row_data):
                table.rows[1].cells[j].text = val

        path = str(tmp_path / "multi_table.docx")
        doc.save(path)
        parser = DocxParser(path)
        result = parser.extract_use_cases()
        assert len(result) == 2

    def test_invalid_file_raises_error(self, tmp_path):
        path = str(tmp_path / "notadoc.docx")
        with open(path, 'w') as f:
            f.write("not a real docx file")
        parser = DocxParser(path)
        with pytest.raises(ValueError):
            parser.extract_use_cases()

    def test_partial_columns(self, tmp_path):
        """Only some columns present — should still work."""
        headers = ["Use Case", "Description", "Résultats Attendus"]
        rows = [["UC001", "Test partiel", "OK"]]
        path = _create_test_docx(tmp_path, headers, rows)
        parser = DocxParser(path)
        result = parser.extract_use_cases()
        assert len(result) == 1
        assert result[0]['use_case_id'] == 'UC001'
        assert result[0]['steps'] == ''  # missing column defaults to empty
